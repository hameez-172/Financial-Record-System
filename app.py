import streamlit as st
import pandas as pd
import sqlite3
import re
from datetime import datetime, date
from fpdf import FPDF
import os
import hashlib
import plotly.express as px
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =========================================================================
# DISPLAY / EXPORT COLUMN CONFIG
# (kept in one place so Records / Credit / Debit / Expense / Liabilities
#  sheets all use the same column order + headers everywhere: on-screen
#  table, editor, CSV export and PDF export)
# =========================================================================

RECORD_DISPLAY_COLUMNS = ['id', 'date', 'client', 'equipment', 'specs', 'qty_per_item',
                          'close_deal', 'actual_cost', 'actual_price_per_item',
                          'other_expenses_per_item', 'paid', 'remaining', 'profit',
                          'team_member', 'status']
RECORD_DISPLAY_HEADERS = ['No.', 'Date', 'Client', 'Equipment', 'Specs', 'Qty',
                           'Close Deal', 'Actual Cost', 'Actual Price/Item',
                           'Other Expenses', 'Paid', 'Remaining', 'Profit',
                           'Team Member', 'Status']
RECORD_COL_WIDTHS = [10, 18, 26, 26, 22, 16, 18, 18, 20, 20, 16, 18, 16, 20, 13]  # sums to 277mm

CREDIT_HEADERS = ["Client Name", "Id No", "Total Payment", "Paid by Client", "Remaining from Client"]
CREDIT_COL_WIDTHS = [60, 35, 60, 60, 62]  # sums to 277mm

DEBIT_HEADERS = ["Client Name", "Id No", "Total Payment", "Paid to Client", "Remaining to be paid"]
DEBIT_COL_WIDTHS = [60, 35, 60, 60, 62]  # sums to 277mm

EXPENSE_HEADERS = ["Description", "Amount"]
EXPENSE_COL_WIDTHS = [130, 60]  # sums to 190mm (A4 portrait usable width)

LIABILITY_HEADERS = ["Date", "Description", "Total Amount", "Paid Amount", "Remaining"]
LIABILITY_COL_WIDTHS = [25, 75, 30, 30, 30]  # sums to 190mm (A4 portrait usable width)

# Doc-number prefix per document type (Point 1: Quotation should print
# "QUO-..." instead of "INV-...", Delivery Challan prints "DC-...";
# the underlying invoice_no stored in the DB never changes, only what
# gets printed on the document).
DOC_NUMBER_PREFIX = {"Invoice": "INV", "Quotation": "QUO", "Delivery Challan": "DC"}

# Status options for the manual status dropdown on every deal entry.
STATUS_OPTIONS = ["Decline", "Approved", "Paid", "Pending"]


# =========================================================================
# FORMATTING HELPERS
# =========================================================================

def _format_date_ddmmyyyy(value):
    try:
        return pd.to_datetime(value).strftime("%d-%m-%Y")
    except Exception:
        return value


def _fmt_money(value):
    try:
        return f"{float(value):,.0f}"
    except (TypeError, ValueError):
        return value


def _prepare_export_df(df, date_cols=(), money_cols=()):
    out = df.copy()
    for col in date_cols:
        if col in out.columns:
            out[col] = out[col].apply(_format_date_ddmmyyyy)
    for col in money_cols:
        if col in out.columns:
            out[col] = out[col].apply(_fmt_money)
    return out


def _parse_csv_floats(s):
    if s is None:
        return []
    out = []
    for p in str(s).split(','):
        p = p.strip()
        if not p or p.lower() in ('none', 'nan'):
            continue
        try:
            out.append(float(p))
        except ValueError:
            continue
    return out


def _sanitize_filename(text):
    """Point 4: strips characters that aren't safe in a filename so the
    downloaded file can be named after the client, e.g. 'ALKHAIR HOSPITAL'."""
    text = str(text) if text is not None else ""
    text = re.sub(r'[\\/*?:"<>|]', '', text).strip()
    return text if text else "Client"


def _doc_type_word(doc_type):
    """Point 4: uppercase word used in the downloaded file name --
    INVOICE / QUOTATION / DELIVERY CHALLAN."""
    return str(doc_type).upper()


def _display_doc_number(invoice_no, doc_type):
    """Point 1: swaps the INV- prefix stored in the DB for the prefix that
    matches what's being printed (QUO- for Quotation, DC- for Delivery
    Challan), without touching the invoice_no actually saved in the DB."""
    prefix = DOC_NUMBER_PREFIX.get(doc_type, "INV")
    invoice_no = str(invoice_no) if invoice_no else ""
    if "-" in invoice_no:
        _, rest = invoice_no.split("-", 1)
        return f"{prefix}-{rest}"
    return f"{prefix}-{invoice_no}" if invoice_no else prefix


_ONES = ["", "ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT", "NINE",
         "TEN", "ELEVEN", "TWELVE", "THIRTEEN", "FOURTEEN", "FIFTEEN", "SIXTEEN",
         "SEVENTEEN", "EIGHTEEN", "NINETEEN"]
_TENS = ["", "", "TWENTY", "THIRTY", "FORTY", "FIFTY", "SIXTY", "SEVENTY", "EIGHTY", "NINETY"]


def _three_digit_words(n):
    words = ""
    if n >= 100:
        words += _ONES[n // 100] + " HUNDRED "
        n %= 100
    if n >= 20:
        words += _TENS[n // 10] + " "
        n %= 10
        if n:
            words += _ONES[n] + " "
    elif n > 0:
        words += _ONES[n] + " "
    return words.strip()


def _number_to_words(n):
    n = int(round(n))
    if n == 0:
        return "ZERO"
    negative = n < 0
    n = abs(n)
    parts = []
    for divisor, name in [(1_000_000_000, "BILLION"), (1_000_000, "MILLION"), (1_000, "THOUSAND")]:
        if n >= divisor:
            parts.append(f"{_three_digit_words(n // divisor)} {name}")
            n %= divisor
    if n > 0:
        parts.append(_three_digit_words(n))
    result = " ".join(parts).strip()
    return ("MINUS " + result) if negative else result


# --- PDF GENERATOR CLASS ---
class InvoicePDF(FPDF):
    def header(self):
        self.set_fill_color(0, 51, 102); self.rect(10, 8, 22, 8, "F")
        blue_w = self.w - 45
        self.set_fill_color(0, 153, 224); self.rect(35, 8, blue_w, 8, "F")
        if os.path.exists("lo.png"): self.image("lo.png", x=10, y=18, w=25)
        self.set_xy(40, 20); self.set_font("Arial", "B", 20); self.set_text_color(20, 40, 80)
        self.cell(0, 10, "Badar Diagnostics & Medical Equipments")

    def footer(self):
        rect_w = self.w - 20
        navy_y = self.h - 37
        blue_y = self.h - 22
        self.set_fill_color(0, 51, 102); self.rect(10, navy_y, rect_w, 15, "F")
        self.set_fill_color(0, 153, 224); self.rect(10, blue_y, rect_w, 8, "F")
        self.set_y(self.h - 35); self.set_text_color(255, 255, 255); self.set_font("Arial", "", 7)
        self.multi_cell(0, 3.5, "Lahore Office: D Block Nawab Town, Lahore   |   Okara Office: Adjacent Ibn-e-Sina Lab, Opposite DHQ, Okara\nPindi Office: Commercial Market, Rawalpindi   |   Bahawalpur Office: Model Town C, Bahawalpur", align="C")
        self.set_y(self.h - 21); self.set_font("Arial", "B", 8)
        self.cell(0, 4, " 0300-7303020, 0334-7303020      E-mail: munir.badar1@gmail.com", align="C")


def _draw_item_table_header(pdf, y):
    pdf.set_xy(25, y)
    pdf.set_draw_color(0, 153, 224)
    pdf.set_font("Arial", "B", 7); pdf.set_fill_color(240, 240, 240)
    headers = ["SR #", "PRODUCT", "SPECS", "QTY", "PRICE PER UNIT IN PKR", "TOTAL PRICE IN PKR"]
    widths = [15, 45, 40, 15, 25, 25]
    _draw_wrapped_row(pdf, headers, widths, line_h=3.3, start_x=25, align="C", fill=True)


def _wrapped_line_count(pdf, text, width):
    """Estimates how many lines a multi_cell(width, ...) call will take for the
    given text -- used both for sizing the Terms & Conditions block AND (now)
    for auto-growing table rows so long names/specs never get cut off, they
    just wrap onto extra lines and the row grows taller automatically."""
    total_lines = 0
    for paragraph in str(text).split("\n"):
        if paragraph.strip() == "":
            total_lines += 1
            continue
        words = paragraph.split(" ")
        current = ""
        for word in words:
            trial = f"{current} {word}".strip()
            if pdf.get_string_width(trial) <= width - 2:
                current = trial
            else:
                total_lines += 1
                current = word
        total_lines += 1
    return max(total_lines, 1)


def _draw_wrapped_row(pdf, values, widths, line_h, start_x, align="C", fill=False):
    """Draws one table row where every cell auto-wraps long text onto multiple
    lines instead of cutting it off. The border for EVERY cell in the row is
    drawn at the same full row height (the tallest cell's height) first, then
    the (possibly multi-line) text is written on top without its own border --
    this keeps the whole row's grid lines aligned instead of only the
    long-text column's box stretching while the rest stay short.
    `align` can be a single alignment for every column, or a list/tuple with
    one alignment per column (e.g. left-align a name column, center the rest).
    `fill` draws a background fill behind the row (used for table headers).
    Text is also centered VERTICALLY within the row's full height, so a
    single-line cell (e.g. a header like "QTY") sits in the middle of its box
    instead of stuck at the top."""
    aligns = list(align) if isinstance(align, (list, tuple)) else [align] * len(values)
    per_cell_lines = [_wrapped_line_count(pdf, v, w) for v, w in zip(values, widths)]
    row_h = max(per_cell_lines) * line_h
    y0 = pdf.get_y()
    x = start_x
    for v, w, a, nl in zip(values, widths, aligns, per_cell_lines):
        pdf.rect(x, y0, w, row_h, "DF" if fill else "D")
        y_text = y0 + (row_h - nl * line_h) / 2.0
        pdf.set_xy(x, y_text)
        pdf.multi_cell(w, line_h, v, border=0, align=a)
        x += w
    pdf.set_xy(start_x, y0 + row_h)
    return row_h


def generate_pdf(deal, items_df, doc_type="Invoice", terms_text=None):
    pdf = InvoicePDF()
    pdf.add_page()
    blue_color = (0, 153, 224)
    pdf.set_draw_color(*blue_color)

    is_challan = (doc_type == "Delivery Challan")

    # 1. Invoice No & Date
    pdf.set_xy(15, 45)
    pdf.set_font("Arial", "B", 12); pdf.set_text_color(*blue_color)
    pdf.cell(10, 5, "No."); pdf.set_text_color(0, 0, 0); pdf.set_font("Arial", "", 12)
    inv_text = _display_doc_number(deal['invoice_no'], doc_type)
    pdf.set_xy(25, 45); pdf.cell(pdf.get_string_width(inv_text), 5, inv_text)
    pdf.line(25, 50, 25 + pdf.get_string_width(inv_text), 50)

    pdf.set_xy(140, 45); pdf.set_font("Arial", "B", 12); pdf.set_text_color(*blue_color)
    pdf.cell(10, 5, "Date"); pdf.set_text_color(0, 0, 0); pdf.set_font("Arial", "", 12)
    date_val = _format_date_ddmmyyyy(deal['date'])
    pdf.set_xy(155, 45); pdf.cell(pdf.get_string_width(date_val), 5, date_val)
    pdf.line(155, 50, 155 + pdf.get_string_width(date_val), 50)

    # 2. Client Name
    pdf.set_text_color(0, 0, 0); pdf.set_xy(15, 58); pdf.set_font("Arial", "B", 12)
    pdf.cell(10, 6, "To: ")
    name_x = pdf.get_x()
    client_name = f"{deal['client']}"
    pdf.set_font("Arial", "B", 12)
    pdf.cell(pdf.get_string_width(client_name), 6, client_name)
    pdf.set_draw_color(0, 0, 0)
    pdf.line(name_x, 64, name_x + pdf.get_string_width(client_name), 64)

    # 3. Table
    pdf.set_xy(0, 70); pdf.set_font("Arial", "B", 16); pdf.cell(210, 8, doc_type.upper(), align="C")

    _draw_item_table_header(pdf, 85)
    pdf.set_font("Arial", "", 9)
    pdf.set_draw_color(0, 153, 224)

    item_widths = [15, 45, 40, 15, 25, 25]
    item_aligns = ["C", "L", "L", "C", "C", "C"]
    line_h = 5

    for i, item in enumerate(items_df.itertuples(), start=1):
        price_txt = "" if is_challan else f"{item.unit_price:,.0f}"
        total_txt = "" if is_challan else f"{item.line_total:,.0f}"
        values = [str(i), str(item.equipment), str(item.specs),
                  f"{item.quantity:g}", price_txt, total_txt]

        n_lines = [_wrapped_line_count(pdf, v, w) for v, w in zip(values, item_widths)]
        row_h = max(n_lines) * line_h

        if pdf.get_y() + row_h > 250:
            pdf.add_page(); _draw_item_table_header(pdf, 45)
            pdf.set_font("Arial", "", 9); pdf.set_draw_color(0, 153, 224)

        _draw_wrapped_row(pdf, values, item_widths, line_h, start_x=25, align=item_aligns)

    if pdf.get_y() + 20 > 250:
        pdf.add_page()

    if not is_challan:
        pdf.set_x(125); pdf.set_font("Arial", "B", 10)
        pdf.cell(40, 8, "Grand Total", 1, 0, "C", True)
        pdf.cell(25, 8, f"{deal['close_deal']:,.0f}", 1, 1, "C", True)

        pdf.set_x(15)
        pdf.set_font("Arial", "B", 8)
        pdf.cell(175, 5, f"{_number_to_words(deal['close_deal'])} RUPEES ONLY", align="R")
        pdf.ln(6)

    show_terms = doc_type == "Quotation" and terms_text and terms_text.strip()
    divider_y = 222

    if show_terms:
        if pdf.get_y() + 10 > 195:
            pdf.add_page()
        terms_y = pdf.get_y() + 10
        pdf.set_xy(15, terms_y)
        pdf.set_font("Arial", "B", 11)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 6, "Terms & Conditions:", ln=1)
        pdf.set_x(15)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(90, 5, terms_text)

    if pdf.get_y() + 10 > divider_y:
        pdf.add_page()

    content_y = divider_y + 3

    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, divider_y, 200, divider_y)

    pdf.set_xy(15, content_y)
    pdf.set_font("Arial", "I", 9)
    pdf.cell(90, 5, "Regards,", ln=1)

    pdf.set_x(15)
    pdf.set_font("Arial", "B", 9)
    pdf.cell(90, 5, "Badar Diagnostics & Medical Equipment, Lahore", ln=1)

    pdf.set_x(15)
    pdf.set_font("Arial", "B", 9)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(90, 5, "Account Details:", ln=1)

    pdf.set_x(15)
    pdf.set_font("Arial", "", 8)
    pdf.set_text_color(0, 0, 0)
    pdf.multi_cell(90, 4, "Badar Diagnostics & Medical Equipment\nFaysal Bank\n0155007000005585")

    if os.path.exists("stamp.jpg"):
        pdf.image("stamp.jpg", x=140, y=content_y, w=35)

    # Point 4: file is named after the client + doc type, e.g.
    # "ALKHAIR HOSPITAL INVOICE.pdf" / "ALKHAIR HOSPITAL QUOTATION.pdf"
    file_path = f"{_sanitize_filename(deal['client'])} {_doc_type_word(doc_type)}.pdf"
    pdf.output(file_path)
    return file_path


# =========================================================================
# WORD (.docx) GENERATOR -- mirrors generate_pdf() so Invoice / Quotation /
# Delivery Challan can be downloaded as an editable Word document too.
# Point 2: header banner + logo, blue-bordered item table, grand total box,
# footer bands and stamp image are all rebuilt here to match the PDF look.
# =========================================================================

_NAVY = RGBColor(0, 51, 102)
_BLUE = RGBColor(0, 153, 224)
_BLACK = RGBColor(0, 0, 0)
_GREY = RGBColor(240, 240, 240)


def _docx_shade_cell(cell, rgb_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tc_pr.append(shd)


def _docx_set_cell_border(cell, color_hex="0099E0", sz=4):
    """Gives a cell a colored border on all four sides (used to mirror the
    PDF's blue table borders, which python-docx's default 'Table Grid'
    style doesn't give us)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color_hex)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _docx_set_col_widths(table, widths_mm):
    """python-docx needs width set on both the column AND every cell for a
    fixed layout to actually stick."""
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_mm):
            cell.width = Mm(w)
    for i, w in enumerate(widths_mm):
        if i < len(table.columns):
            table.columns[i].width = Mm(w)


def _docx_set_cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _docx_header_banner(doc):
    """Mirrors InvoicePDF.header(): a thin navy+blue strip, then the logo
    next to the company name in navy text -- same as the PDF."""
    strip = doc.add_table(rows=1, cols=2)
    _docx_set_col_widths(strip, [22, 158])
    _docx_shade_cell(strip.cell(0, 0), "003366")
    _docx_shade_cell(strip.cell(0, 1), "0099E0")
    for cell in strip.rows[0].cells:
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(" ")
        run.font.size = Pt(2)

    header_table = doc.add_table(rows=1, cols=2)
    _docx_set_col_widths(header_table, [30, 150])
    logo_cell = header_table.cell(0, 0)
    logo_cell.text = ""
    if os.path.exists("lo.png"):
        run = logo_cell.paragraphs[0].add_run()
        try:
            run.add_picture("lo.png", width=Mm(22))
        except Exception:
            pass
    name_cell = header_table.cell(0, 1)
    name_cell.text = ""
    p = name_cell.paragraphs[0]
    run = p.add_run("Badar Diagnostics & Medical Equipments")
    run.font.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(20, 40, 80)


def _docx_footer_bands(doc):
    """Mirrors InvoicePDF.footer(): navy band with the office addresses,
    then a blue band with the phone/email, both in white text."""
    footer_table = doc.add_table(rows=2, cols=1)
    footer_table.autofit = True

    navy_cell = footer_table.cell(0, 0)
    _docx_shade_cell(navy_cell, "003366")
    navy_cell.text = ""
    p1 = navy_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("Lahore Office: D Block Nawab Town, Lahore   |   Okara Office: Adjacent Ibn-e-Sina Lab, Opposite DHQ, Okara")
    r1.font.size = Pt(7); r1.font.color.rgb = RGBColor(255, 255, 255)
    p2 = navy_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Pindi Office: Commercial Market, Rawalpindi   |   Bahawalpur Office: Model Town C, Bahawalpur")
    r2.font.size = Pt(7); r2.font.color.rgb = RGBColor(255, 255, 255)

    blue_cell = footer_table.cell(1, 0)
    _docx_shade_cell(blue_cell, "0099E0")
    blue_cell.text = ""
    p3 = blue_cell.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("0300-7303020, 0334-7303020      E-mail: munir.badar1@gmail.com")
    r3.font.bold = True; r3.font.size = Pt(8); r3.font.color.rgb = RGBColor(255, 255, 255)


def _docx_signature_block(doc):
    """Mirrors the PDF's Regards / Account Details block, with the stamp
    image placed on the right the same way pdf.image() places it in
    generate_pdf()."""
    doc.add_paragraph("_" * 90)
    sig_table = doc.add_table(rows=1, cols=2)
    _docx_set_col_widths(sig_table, [120, 60])

    left = sig_table.cell(0, 0)
    left.text = ""
    p = left.paragraphs[0]
    r = p.add_run("Regards,")
    r.font.italic = True; r.font.size = Pt(9)
    p2 = left.add_paragraph()
    r2 = p2.add_run("Badar Diagnostics & Medical Equipment, Lahore")
    r2.font.bold = True; r2.font.size = Pt(9)
    p3 = left.add_paragraph()
    r3 = p3.add_run("Account Details:")
    r3.font.bold = True; r3.font.size = Pt(9); r3.font.color.rgb = _NAVY
    p4 = left.add_paragraph()
    r4 = p4.add_run("Badar Diagnostics & Medical Equipment\nFaysal Bank\n0155007000005585")
    r4.font.size = Pt(8)

    right = sig_table.cell(0, 1)
    right.text = ""
    if os.path.exists("stamp.jpg"):
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = rp.add_run()
        try:
            run.add_picture("stamp.jpg", width=Mm(35))
        except Exception:
            pass


def generate_docx(deal, items_df, doc_type="Invoice", terms_text=None):
    is_challan = (doc_type == "Delivery Challan")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    _docx_header_banner(doc)
    doc.add_paragraph()

    # No. / Date row
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.autofit = True
    left_p = meta_table.cell(0, 0).paragraphs[0]
    r = left_p.add_run("No.  "); r.font.bold = True; r.font.color.rgb = _BLUE
    r2 = left_p.add_run(_display_doc_number(deal['invoice_no'], doc_type)); r2.font.bold = False

    right_p = meta_table.cell(0, 1).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = right_p.add_run("Date  "); r3.font.bold = True; r3.font.color.rgb = _BLUE
    r4 = right_p.add_run(_format_date_ddmmyyyy(deal['date']))

    doc.add_paragraph()

    # Client
    client_p = doc.add_paragraph()
    r5 = client_p.add_run("To: "); r5.font.bold = True
    r6 = client_p.add_run(f"{deal['client']}"); r6.font.bold = True

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(doc_type.upper())
    title_run.font.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph()

    # Items table
    headers = ["SR #", "PRODUCT", "SPECS", "QTY", "PRICE PER UNIT IN PKR", "TOTAL PRICE IN PKR"]
    item_table = doc.add_table(rows=1, cols=len(headers))
    item_table.style = "Table Grid"
    item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = item_table.cell(0, i)
        _docx_shade_cell(cell, "F0F0F0")
        _docx_set_cell_text(cell, h, bold=True, size=8)

    for i, item in enumerate(items_df.itertuples(), start=1):
        row_cells = item_table.add_row().cells
        price_txt = "" if is_challan else f"{item.unit_price:,.0f}"
        total_txt = "" if is_challan else f"{item.line_total:,.0f}"
        values = [str(i), str(item.equipment), str(item.specs), f"{item.quantity:g}", price_txt, total_txt]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
        for cell, val, al in zip(row_cells, values, aligns):
            _docx_set_cell_text(cell, val, size=9, align=al)

    # Blue borders on every cell of the item table, to match the PDF's
    # blue-bordered grid (header row included).
    for row in item_table.rows:
        for cell in row.cells:
            _docx_set_cell_border(cell, "0099E0", sz=4)

    doc.add_paragraph()

    if not is_challan:
        total_table = doc.add_table(rows=1, cols=2)
        total_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        _docx_set_col_widths(total_table, [30, 25])
        for cell in total_table.rows[0].cells:
            _docx_shade_cell(cell, "F0F0F0")
            _docx_set_cell_border(cell, "0099E0", sz=6)
        _docx_set_cell_text(total_table.cell(0, 0), "Grand Total", bold=True, size=10)
        _docx_set_cell_text(total_table.cell(0, 1), f"{deal['close_deal']:,.0f}", bold=True, size=10)

        words_p = doc.add_paragraph()
        words_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        words_run = words_p.add_run(f"{_number_to_words(deal['close_deal'])} RUPEES ONLY")
        words_run.font.bold = True
        words_run.font.size = Pt(8)

    if doc_type == "Quotation" and terms_text and terms_text.strip():
        doc.add_paragraph()
        terms_head = doc.add_paragraph()
        r7 = terms_head.add_run("Terms & Conditions:")
        r7.font.bold = True
        r7.font.size = Pt(11)
        for line in terms_text.split("\n"):
            if line.strip():
                terms_p = doc.add_paragraph()
                terms_p.add_run(line).font.size = Pt(10)

    doc.add_paragraph()
    _docx_signature_block(doc)
    doc.add_paragraph()
    _docx_footer_bands(doc)

    # Point 4: file is named after the client + doc type, e.g.
    # "ALKHAIR HOSPITAL INVOICE.docx" / "ALKHAIR HOSPITAL QUOTATION.docx"
    file_path = f"{_sanitize_filename(deal['client'])} {_doc_type_word(doc_type)}.docx"
    doc.save(file_path)
    return file_path


# =========================================================================
# GENERIC SHEET EXPORT (Records / Credit Sheet / Debit Sheet / Expense Sheet
# / Liabilities Sheet)
# =========================================================================

def _draw_sheet_table_header(pdf, headers, col_widths, y, start_x=10):
    pdf.set_xy(start_x, y)
    pdf.set_draw_color(0, 153, 224)
    pdf.set_font("Arial", "B", 8)
    pdf.set_fill_color(240, 240, 240)
    pdf.set_text_color(0, 0, 0)
    for h, w in zip(headers, col_widths):
        pdf.cell(w, 8, str(h), 1, 0, "C", True)


def generate_sheet_pdf(df, headers, col_widths, title, filename_prefix, orientation="L"):
    pdf = InvoicePDF(orientation=orientation)
    pdf.add_page()
    start_x = 10
    title_y = 30 if orientation == "L" else 45
    pdf.set_xy(0, title_y)
    pdf.set_font("Arial", "B", 16); pdf.set_text_color(0, 0, 0)
    pdf.cell(pdf.w, 8, title.upper(), align="C")

    table_y = title_y + 14
    _draw_sheet_table_header(pdf, headers, col_widths, table_y, start_x)
    pdf.set_font("Arial", "", 8); pdf.set_text_color(0, 0, 0); pdf.set_draw_color(0, 153, 224)

    line_h = 5
    y = table_y + 8
    bottom_limit = pdf.h - 45  # leave room for the letterhead footer band

    for _, row in df.iterrows():
        values = ["" if pd.isna(v) else str(v) for v in row.tolist()]
        n_lines = [_wrapped_line_count(pdf, v, w) for v, w in zip(values, col_widths)]
        row_h = max(n_lines) * line_h

        if y + row_h > bottom_limit:
            pdf.add_page()
            y = title_y
            _draw_sheet_table_header(pdf, headers, col_widths, y, start_x)
            pdf.set_font("Arial", "", 8); pdf.set_draw_color(0, 153, 224)
            y += 8

        pdf.set_xy(start_x, y)
        _draw_wrapped_row(pdf, values, col_widths, line_h, start_x)
        y += row_h

    path = f"{filename_prefix}.pdf"
    pdf.output(path)
    return path


def _df_to_csv_bytes(df):
    return df.to_csv(index=False).encode("utf-8-sig")


def _apply_date_filter(df, date_col, start, end):
    if df is None or df.empty or date_col not in df.columns or start is None or end is None:
        return df
    parsed = pd.to_datetime(df[date_col], errors='coerce')
    mask = (parsed.dt.date >= start) & (parsed.dt.date <= end) | parsed.isna()
    return df[mask]


# =========================================================================
# DATABASE CONNECTION (Turso, via plain HTTP -- no Rust/native build needed,
# so it works on any Streamlit Cloud Python version). Falls back to local
# sqlite3 if Turso secrets aren't configured (e.g. local dev without a
# .streamlit/secrets.toml).
# =========================================================================

import requests
import base64


class _TursoCursor:
    """Mimics just enough of the sqlite3 Cursor interface (execute,
    fetchall, fetchone, description, lastrowid) for this app's needs."""

    def __init__(self, conn):
        self._conn = conn
        self._rows = []
        self.description = None
        self.lastrowid = None
        self.rowcount = -1

    def execute(self, sql, params=None):
        result = self._conn._run(sql, params or ())
        self._rows = result["rows"]
        self.description = result["description"]
        self.lastrowid = result["last_insert_rowid"]
        self.rowcount = result["affected_row_count"]
        return self

    def fetchall(self):
        return self._rows

    def fetchone(self):
        return self._rows[0] if self._rows else None

    def close(self):
        pass


class TursoHTTPConnection:
    """Talks to a Turso database over its plain HTTP pipeline API
    (https://docs.turso.tech/sdk/http/reference). No native/Rust
    dependency at all -- just the `requests` library -- so it always
    installs cleanly regardless of the Python version Streamlit Cloud
    happens to be running."""

    def __init__(self, url, token):
        http_url = url.replace("libsql://", "https://").rstrip("/")
        self._pipeline_url = http_url + "/v2/pipeline"
        self._token = token
        self._session = requests.Session()

    @staticmethod
    def _to_arg(v):
        if v is None:
            return {"type": "null"}
        if isinstance(v, bool):
            return {"type": "integer", "value": str(int(v))}
        if isinstance(v, int):
            return {"type": "integer", "value": str(v)}
        if isinstance(v, float):
            return {"type": "float", "value": v}
        if isinstance(v, (bytes, bytearray)):
            return {"type": "blob", "base64": base64.b64encode(v).decode("ascii")}
        return {"type": "text", "value": str(v)}

    @staticmethod
    def _from_cell(cell):
        if cell is None:
            return None
        t = cell.get("type")
        if t == "null":
            return None
        if t == "integer":
            return int(cell["value"])
        if t == "float":
            return float(cell["value"])
        if t == "text":
            return cell["value"]
        if t == "blob":
            return base64.b64decode(cell["base64"])
        return cell.get("value")

    def _run(self, sql, params):
        args = [self._to_arg(p) for p in params]
        payload = {"requests": [
            {"type": "execute", "stmt": {"sql": sql, "args": args}},
            {"type": "close"},
        ]}
        headers = {"Authorization": f"Bearer {self._token}", "Content-Type": "application/json"}
        resp = self._session.post(self._pipeline_url, json=payload, headers=headers, timeout=30)
        resp.raise_for_status()
        data = resp.json()
        results = data.get("results", [])
        if not results:
            raise RuntimeError("Turso: empty response received.")
        first = results[0]
        if first.get("type") == "error":
            err = first.get("error", {})
            raise RuntimeError(f"Turso SQL error: {err.get('message', err)}")
        exec_result = first["response"]["result"]
        cols = [c["name"] for c in exec_result.get("cols", [])]
        rows = [tuple(self._from_cell(cell) for cell in row) for row in exec_result.get("rows", [])]
        last_id = exec_result.get("last_insert_rowid")
        last_id = int(last_id) if last_id is not None else None
        return {
            "rows": rows,
            "description": [(c, None, None, None, None, None, None) for c in cols],
            "last_insert_rowid": last_id,
            "affected_row_count": exec_result.get("affected_row_count", 0),
        }

    def cursor(self):
        return _TursoCursor(self)

    def execute(self, sql, params=None):
        cur = self.cursor()
        cur.execute(sql, params)
        return cur

    def commit(self):
        pass  # every statement is committed immediately over HTTP

    def sync(self):
        pass  # kept only so existing `if hasattr(conn, "sync")` calls stay harmless

    def close(self):
        pass


def read_sql_df(query, conn, params=None):
    """Drop-in replacement for read_sql_df() that works identically for both
    sqlite3.Connection and TursoHTTPConnection (pandas' own read_sql only
    recognizes real sqlite3.Connection objects)."""
    cur = conn.cursor()
    cur.execute(query, params or ())
    cols = [d[0] for d in cur.description] if cur.description else []
    rows = cur.fetchall()
    return pd.DataFrame(rows, columns=cols)


def get_connection():
    """Returns a database connection. If Turso secrets are configured, it
    connects to the Turso cloud database over plain HTTP, so data survives
    Streamlit Cloud reboots/redeploys. Falls back to local sqlite3 if
    secrets aren't set (e.g. local dev without a .streamlit/secrets.toml)."""
    try:
        url = st.secrets.get("TURSO_DATABASE_URL")
        token = st.secrets.get("TURSO_AUTH_TOKEN")
    except Exception:
        url = None
        token = None

    if url and token:
        return TursoHTTPConnection(url, token)

    return sqlite3.connect('enterprise.db')


# --- APP SETUP ---
def init_db():
    conn = get_connection()
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS business_deals
                  (id INTEGER PRIMARY KEY AUTOINCREMENT, date TEXT, invoice_no TEXT, client TEXT,
                  equipment TEXT, specs TEXT, qty_per_item TEXT, close_deal REAL, actual_cost REAL,
                  actual_price_per_item TEXT, other_expenses_per_item TEXT,
                  paid REAL, remaining REAL, profit REAL, team_member TEXT, status TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS deal_items
                  (id INTEGER PRIMARY KEY AUTOINCREMENT, deal_id INTEGER, equipment TEXT, specs TEXT,
                  quantity REAL, unit_price REAL, unit_actual_cost REAL, other_expenses REAL,
                  line_total REAL, line_actual_cost REAL,
                  FOREIGN KEY(deal_id) REFERENCES business_deals(id))''')
    c.execute('''CREATE TABLE IF NOT EXISTS credit_manual
                  (id INTEGER PRIMARY KEY AUTOINCREMENT, client TEXT, total_payment REAL,
                  paid_by_client REAL, remaining_from_client REAL)''')
    c.execute('''CREATE TABLE IF NOT EXISTS debit_manual
                  (id INTEGER PRIMARY KEY AUTOINCREMENT, client TEXT, total_payment REAL,
                  paid_to_client REAL, remaining_to_be_paid REAL)''')
    c.execute('''CREATE TABLE IF NOT EXISTS daily_expenses
                  (id INTEGER PRIMARY KEY AUTOINCREMENT, date TEXT, category TEXT,
                  description TEXT, amount REAL)''')
    # Point 3: new table backing the Liabilities tab.
    c.execute('''CREATE TABLE IF NOT EXISTS liabilities
                  (id INTEGER PRIMARY KEY AUTOINCREMENT, date TEXT, description TEXT,
                  total_amount REAL, paid_amount REAL, remaining REAL)''')

    existing_cols = [row[1] for row in c.execute("PRAGMA table_info(business_deals)").fetchall()]
    if 'equipment' not in existing_cols:
        c.execute("ALTER TABLE business_deals ADD COLUMN equipment TEXT")
    if 'specs' not in existing_cols:
        c.execute("ALTER TABLE business_deals ADD COLUMN specs TEXT")
    if 'actual_price_per_item' not in existing_cols:
        c.execute("ALTER TABLE business_deals ADD COLUMN actual_price_per_item TEXT")
    if 'qty_per_item' not in existing_cols:
        c.execute("ALTER TABLE business_deals ADD COLUMN qty_per_item TEXT")
    if 'other_expenses_per_item' not in existing_cols:
        c.execute("ALTER TABLE business_deals ADD COLUMN other_expenses_per_item TEXT")

    existing_item_cols = [row[1] for row in c.execute("PRAGMA table_info(deal_items)").fetchall()]
    if 'other_expenses' not in existing_item_cols:
        c.execute("ALTER TABLE deal_items ADD COLUMN other_expenses REAL DEFAULT 0")

    conn.commit()
    if hasattr(conn, "sync"):
        conn.sync()
    conn.close()


init_db()
st.set_page_config(page_title="Badar Diagnostics & Medical Equipments", layout="wide")

# =========================================================================
# MOBILE-FRIENDLY STYLING -- shrinks paddings, fonts, buttons, tabs and
# tables on small screens so the app is usable on a phone without the
# desktop layout looking oversized. EVERYTHING inside the
# @media (max-width: 768px) block below only fires on mobile-sized
# viewports -- the desktop/laptop layout above/outside it is untouched.
# =========================================================================
st.markdown("""
<style>
/* Tighten the default page padding everywhere */
.block-container {
    padding-top: 2.2rem;
    padding-bottom: 2rem;
    padding-left: 1rem;
    padding-right: 1rem;
}

/* Push the tab bar down a bit and give tab labels more breathing room so
   they don't get clipped by the top toolbar / page edge. */
div[data-testid="stTabs"] {
    margin-top: 0.6rem;
}
div[data-baseweb="tab-list"] {
    margin-top: 6px !important;
    padding-top: 6px !important;
    padding-bottom: 2px !important;
}
button[data-baseweb="tab"] {
    padding-top: 10px !important;
    padding-bottom: 10px !important;
}

@media (max-width: 768px) {
    .block-container {
        padding-top: 1.2rem;
        padding-left: 0.4rem;
        padding-right: 0.4rem;
        padding-bottom: 0.8rem;
    }

    /* Titles / headers -- more compact so more fits on screen at once */
    h1 { font-size: 1.2rem !important; margin-bottom: 0.3rem !important; }
    h2 { font-size: 1.05rem !important; margin-bottom: 0.25rem !important; }
    h3 { font-size: 0.95rem !important; margin-bottom: 0.2rem !important; }
    .stMarkdown p { font-size: 0.82rem !important; }

    /* Tabs -- smaller, scrollable, tighter padding, but still clear of the top edge */
    button[data-baseweb="tab"] {
        font-size: 0.72rem !important;
        padding: 0.45rem 0.4rem !important;
    }
    div[data-baseweb="tab-list"] {
        margin-top: 8px !important;
        gap: 1px !important;
        overflow-x: auto !important;
        flex-wrap: nowrap !important;
    }

    /* Buttons -- smaller and tighter so rows of action buttons don't wrap
       and eat vertical space */
    .stButton button, .stDownloadButton button, .stFormSubmitButton button {
        font-size: 0.74rem !important;
        padding: 0.28rem 0.5rem !important;
        white-space: normal !important;
        min-height: 2rem !important;
    }

    /* Text inputs / number inputs / selects -- more compact */
    .stTextInput input, .stNumberInput input, .stSelectbox div, .stDateInput input {
        font-size: 0.78rem !important;
        padding: 0.25rem 0.4rem !important;
    }
    label, .stTextInput label, .stNumberInput label, .stSelectbox label, .stDateInput label {
        font-size: 0.72rem !important;
        margin-bottom: 0.1rem !important;
    }

    /* Metrics -- shrink so the 4-metric rows fit without excessive wrapping */
    div[data-testid="stMetricValue"] { font-size: 1rem !important; }
    div[data-testid="stMetricLabel"] { font-size: 0.68rem !important; }
    div[data-testid="stMetric"] { padding: 0.2rem !important; }

    /* Dataframes / data editors -- allow horizontal scroll instead of squeezing,
       and shrink row/cell font so more of the table is readable at once */
    div[data-testid="stDataFrame"], div[data-testid="stDataEditor"] {
        font-size: 0.68rem !important;
    }

    /* Captions -- smaller, tighter line height */
    .stCaption, [data-testid="stCaptionContainer"] {
        font-size: 0.68rem !important;
        line-height: 1.2 !important;
    }

    /* Containers / cards -- compact padding so nested sections don't feel
       spread out */
    div[data-testid="stVerticalBlockBorderWrapper"] {
        padding: 0.3rem !important;
    }

    /* Reduce the default vertical gap Streamlit inserts between stacked
       elements/columns so mobile screens feel dense rather than sparse */
    div[data-testid="stVerticalBlock"] {
        gap: 0.35rem !important;
    }
    div[data-testid="stHorizontalBlock"] {
        gap: 0.3rem !important;
    }
    div[data-testid="column"] {
        padding: 0 0.15rem !important;
    }

    /* Dividers -- tighter spacing */
    hr {
        margin-top: 0.5rem !important;
        margin-bottom: 0.5rem !important;
    }

    /* Sidebar narrower on mobile */
    section[data-testid="stSidebar"] { min-width: 220px !important; }
}
</style>
""", unsafe_allow_html=True)


# =========================================================================
# LOGIN GATE -- no username/password, no access to the app. Credentials are
# best kept in .streamlit/secrets.toml (separate from the code), but if
# secrets.toml is not found, the default username/password below are used
# -- BE SURE TO CHANGE THESE.
# =========================================================================
_DEFAULT_USERNAME = "admin"
_DEFAULT_PASSWORD = "changeme123"  # <-- change this immediately after first login


def _hash_password(password: str) -> str:
    return hashlib.sha256(str(password).encode("utf-8")).hexdigest()


def _get_credentials():
    try:
        username = st.secrets.get("APP_USERNAME", _DEFAULT_USERNAME)
    except Exception:
        username = _DEFAULT_USERNAME
    try:
        password_hash = st.secrets.get("APP_PASSWORD_HASH", None)
    except Exception:
        password_hash = None
    if not password_hash:
        password_hash = _hash_password(_DEFAULT_PASSWORD)
    return username, password_hash


def _login_gate():
    if st.session_state.get("authenticated"):
        return True

    st.markdown(
        "<div style='max-width:420px;margin:80px auto 0 auto;'>",
        unsafe_allow_html=True,
    )
    st.markdown("## 🔒 Badar Diagnostics & Medical Equipments")
    st.caption("Please log in to continue.")
    with st.form("login_form"):
        entered_username = st.text_input("Username")
        entered_password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Login", use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

    if submitted:
        correct_username, correct_password_hash = _get_credentials()
        if entered_username == correct_username and _hash_password(entered_password) == correct_password_hash:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Invalid username or password.")
    return False


if not _login_gate():
    st.stop()

with st.sidebar:
    st.header("⚙️ Account")
    if st.button("🚪 Logout"):
        st.session_state.authenticated = False
        st.rerun()

if 'business_df' not in st.session_state:
    conn = get_connection()
    st.session_state.business_df = read_sql_df("SELECT * FROM business_deals", conn); conn.close()

if 'temp_items' not in st.session_state:
    st.session_state.temp_items = []

if 'update_temp_items' not in st.session_state:
    st.session_state.update_temp_items = []

if 'editing_deal_id' not in st.session_state:
    st.session_state.editing_deal_id = None

if 'confirm_delete_deal_id' not in st.session_state:
    st.session_state.confirm_delete_deal_id = None

if 'credit_manual_df' not in st.session_state or 'debit_manual_df' not in st.session_state:
    conn = get_connection()
    st.session_state.credit_manual_df = read_sql_df("SELECT * FROM credit_manual", conn)
    st.session_state.debit_manual_df = read_sql_df("SELECT * FROM debit_manual", conn)
    conn.close()

if 'expense_df' not in st.session_state:
    conn = get_connection()
    st.session_state.expense_df = read_sql_df("SELECT * FROM daily_expenses", conn)
    conn.close()

if 'liability_df' not in st.session_state:
    conn = get_connection()
    st.session_state.liability_df = read_sql_df("SELECT * FROM liabilities", conn)
    conn.close()

# Point 3: current tabs are Home Finance, Business Deals, Credit/Debit/Expense
# Sheets, Analytics, Liabilities -- plus the new dedicated "Approved" tab
# where entries whose Status is set to "Approved" are grouped together.
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "🏠 Home Finance", "💼 Business Deals", "💳 Credit/Debit/Expense Sheets",
    "📊 Analytics", "📉 Liabilities", "✅ Approved"
])


def _set_deal_status_cb(deal_id, new_status):
    """Shared helper: updates a single deal's status column directly (used by
    the Approved tab's status dropdown and its 'Mark as Paid' payment action).
    Keeps paid/remaining untouched unless the new status is 'Paid', in which
    case the deal is treated as fully settled (paid = close_deal, remaining = 0)
    -- this is the 'payment processed' step for an approved entry."""
    deal_id = int(deal_id)
    conn = get_connection()
    cur = conn.cursor()
    if new_status == "Paid":
        row = st.session_state.business_df[st.session_state.business_df['id'] == deal_id].iloc[0]
        close_deal = row['close_deal']
        cur.execute("UPDATE business_deals SET status=?, paid=?, remaining=? WHERE id=?",
                    (new_status, close_deal, 0.0, deal_id))
    else:
        cur.execute("UPDATE business_deals SET status=? WHERE id=?", (new_status, deal_id))
    conn.commit()
    if hasattr(conn, "sync"):
        conn.sync()
    st.session_state.business_df = read_sql_df("SELECT * FROM business_deals", conn)
    conn.close()


def _highlight_paid_status_cell(val):
    """Highlights ONLY the Status cell (never the whole row) when its value
    is 'Paid'. Subtle, high-contrast green so the text stays readable."""
    if str(val).strip().lower() == 'paid':
        return 'background-color: #c6f6d5; color: #14532d; font-weight: 600;'
    return ''


# ---------------- TAB 2: BUSINESS DEALS ----------------
with tab2:
    st.title("💼 Business Deals")

    def _add_item_cb():
        name = st.session_state.item_name_input
        if name and name.strip():
            qty = st.session_state.item_qty_input
            price = st.session_state.item_price_input
            cost = st.session_state.item_cost_input
            other = st.session_state.item_other_input
            st.session_state.temp_items.append({
                'equipment': name,
                'specs': st.session_state.item_specs_input,
                'quantity': qty,
                'unit_price': price,
                'unit_actual_cost': cost,
                'other_expenses': other,
                'line_total': qty * price,
                'line_actual_cost': qty * cost + other,
            })
            st.session_state.item_name_input = ""
            st.session_state.item_specs_input = ""
            st.session_state.item_qty_input = 1
            st.session_state.item_price_input = 0.0
            st.session_state.item_cost_input = 0.0
            st.session_state.item_other_input = 0.0
            st.session_state.add_item_warning = False
        else:
            st.session_state.add_item_warning = True

    def _remove_item_cb(idx):
        if 0 <= idx < len(st.session_state.temp_items):
            st.session_state.temp_items.pop(idx)

    def _log_deal_cb():
        if not st.session_state.temp_items:
            st.session_state.deal_message = ("error", "Please add at least one product first.")
            return
        if not st.session_state.deal_client.strip():
            st.session_state.deal_message = ("error", "Client Name is required.")
            return

        items = st.session_state.temp_items
        close_deal = sum(i['line_total'] for i in items)
        actual_cost = sum(i['line_actual_cost'] for i in items)
        paid = st.session_state.deal_paid
        remaining = close_deal - paid
        profit = close_deal - actual_cost
        status = "Paid" if remaining <= 0 else "Pending"
        inv_no = f"INV-{datetime.now().strftime('%Y%m%d%H%M%S')}"

        if len(items) == 1:
            equipment_display = items[0]['equipment']
            specs_display = items[0]['specs']
        else:
            equipment_display = ", ".join(i['equipment'] for i in items)
            specs_display = "Multiple Items"

        qty_display = ", ".join(f"{i['quantity']:g}" for i in items)
        actual_price_display = ", ".join(f"{i['unit_actual_cost']:.0f}" for i in items)
        other_expenses_display = ", ".join(f"{i['other_expenses']:.0f}" for i in items)

        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""INSERT INTO business_deals
            (date, invoice_no, client, equipment, specs, qty_per_item, close_deal, actual_cost,
             actual_price_per_item, other_expenses_per_item, paid, remaining, profit, team_member, status)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (datetime.now().strftime("%Y-%m-%d"), inv_no, st.session_state.deal_client, equipment_display,
             specs_display, qty_display, close_deal, actual_cost, actual_price_display, other_expenses_display,
             paid, remaining, profit, st.session_state.deal_team_member, status))
        deal_id = cur.lastrowid

        for item in items:
            cur.execute("""INSERT INTO deal_items
                (deal_id, equipment, specs, quantity, unit_price, unit_actual_cost, other_expenses, line_total, line_actual_cost)
                VALUES (?,?,?,?,?,?,?,?,?)""",
                (deal_id, item['equipment'], item['specs'], item['quantity'], item['unit_price'],
                 item['unit_actual_cost'], item['other_expenses'], item['line_total'], item['line_actual_cost']))

        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()
        st.session_state.business_df = read_sql_df("SELECT * FROM business_deals", conn)
        conn.close()

        st.session_state.temp_items = []
        st.session_state.deal_client = ""
        st.session_state.deal_team_member = ""
        st.session_state.deal_paid = 0.0
        st.session_state.deal_message = ("success", f"Deal {inv_no} saved successfully!")

    def _save_records_edits(edited):
        conn = get_connection()
        cur = conn.cursor()
        for row in edited.to_dict("records"):
            deal_id = row.get('id')
            close_deal = row.get('close_deal', 0) or 0
            paid = row.get('paid', 0) or 0

            qty_list = _parse_csv_floats(row.get('qty_per_item'))
            price_list = _parse_csv_floats(row.get('actual_price_per_item'))
            other_list = _parse_csv_floats(row.get('other_expenses_per_item'))

            if qty_list and price_list and other_list and len(qty_list) == len(price_list) == len(other_list):
                line_costs = [q * p + o for q, p, o in zip(qty_list, price_list, other_list)]
                actual_cost = sum(line_costs)
                item_ids = [r[0] for r in cur.execute(
                    "SELECT id FROM deal_items WHERE deal_id=? ORDER BY id", (int(deal_id),)).fetchall()]
                if len(item_ids) == len(qty_list):
                    for item_id, q, p, o, lc in zip(item_ids, qty_list, price_list, other_list, line_costs):
                        cur.execute("UPDATE deal_items SET quantity=?, unit_actual_cost=?, other_expenses=?, "
                                    "line_actual_cost=? WHERE id=?", (q, p, o, lc, item_id))
            else:
                actual_cost = row.get('actual_cost', 0) or 0

            remaining_edited = row.get('remaining', None)
            if remaining_edited is not None and abs(remaining_edited - (close_deal - paid)) > 0.01:
                remaining = remaining_edited
                paid = close_deal - remaining
            else:
                remaining = close_deal - paid
            profit = close_deal - actual_cost
            # Status is now a manually-controlled dropdown (Decline / Approved /
            # Paid / Pending) rather than auto-derived from Remaining, so we take
            # whatever value the user picked in the editor and keep it as-is.
            status = row.get('status') if row.get('status') in STATUS_OPTIONS else (
                "Paid" if remaining <= 0.01 else "Pending")
            cur.execute("""UPDATE business_deals SET date=?, client=?, equipment=?, specs=?, qty_per_item=?,
                           close_deal=?, actual_cost=?, actual_price_per_item=?, other_expenses_per_item=?,
                           paid=?, remaining=?, profit=?, team_member=?, status=? WHERE id=?""",
                        (row.get('date'), row.get('client'), row.get('equipment'), row.get('specs'),
                         row.get('qty_per_item'), close_deal, actual_cost, row.get('actual_price_per_item'),
                         row.get('other_expenses_per_item'), paid, remaining, profit, row.get('team_member'),
                         status, deal_id))
        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()
        st.session_state.business_df = read_sql_df("SELECT * FROM business_deals", conn)
        conn.close()

    def _delete_deal_cb(deal_id):
        """Point 4: permanently deletes a whole deal (and its line items)."""
        deal_id = int(deal_id)
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("DELETE FROM deal_items WHERE deal_id=?", (deal_id,))
        cur.execute("DELETE FROM business_deals WHERE id=?", (deal_id,))
        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()
        st.session_state.business_df = read_sql_df("SELECT * FROM business_deals", conn)
        conn.close()
        if st.session_state.get("editing_deal_id") == deal_id:
            st.session_state.editing_deal_id = None
            st.session_state.update_temp_items = []
        st.session_state.confirm_delete_deal_id = None
        st.session_state.deal_delete_message = ("success", f"Deal #{deal_id} permanently deleted.")

    def _request_delete_deal_cb(deal_id):
        st.session_state.confirm_delete_deal_id = int(deal_id)

    def _cancel_delete_deal_cb():
        st.session_state.confirm_delete_deal_id = None

    def _save_existing_items_cb(deal_id, edited_df):
        """Point 3: lets you correct a wrong equipment/specs/qty/price entry on an
        existing deal, and also delete individual line items (via the editor's
        built-in row-delete). Recomputes the parent deal's totals afterwards."""
        deal_id = int(deal_id)
        conn = get_connection()
        cur = conn.cursor()
        kept_ids = []
        for row in edited_df.to_dict("records"):
            if not str(row.get('equipment', '')).strip():
                continue
            qty = row.get('quantity') or 0
            price = row.get('unit_price') or 0
            cost = row.get('unit_actual_cost') or 0
            other = row.get('other_expenses') or 0
            line_total = qty * price
            line_actual_cost = qty * cost + other
            raw_id = row.get('id')
            if raw_id is not None and not pd.isna(raw_id):
                item_id = int(raw_id)
                kept_ids.append(item_id)
                cur.execute("""UPDATE deal_items SET equipment=?, specs=?, quantity=?, unit_price=?,
                               unit_actual_cost=?, other_expenses=?, line_total=?, line_actual_cost=?
                               WHERE id=?""",
                            (row.get('equipment'), row.get('specs'), qty, price, cost, other,
                             line_total, line_actual_cost, item_id))
            else:
                cur.execute("""INSERT INTO deal_items
                    (deal_id, equipment, specs, quantity, unit_price, unit_actual_cost, other_expenses, line_total, line_actual_cost)
                    VALUES (?,?,?,?,?,?,?,?,?)""",
                    (deal_id, row.get('equipment'), row.get('specs'), qty, price, cost, other,
                     line_total, line_actual_cost))

        all_existing_ids = [r[0] for r in cur.execute(
            "SELECT id FROM deal_items WHERE deal_id=?", (deal_id,)).fetchall()]
        for item_id in all_existing_ids:
            if item_id not in kept_ids:
                cur.execute("DELETE FROM deal_items WHERE id=?", (item_id,))
        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()

        all_items_df = read_sql_df("SELECT * FROM deal_items WHERE deal_id=?", conn, params=(deal_id,))
        if all_items_df.empty:
            close_deal = 0.0; actual_cost = 0.0
            equipment_display = ""; specs_display = ""
            qty_display = ""; actual_price_display = ""; other_expenses_display = ""
        else:
            close_deal = all_items_df['line_total'].sum()
            actual_cost = all_items_df['line_actual_cost'].sum()
            if len(all_items_df) == 1:
                equipment_display = all_items_df.iloc[0]['equipment']
                specs_display = all_items_df.iloc[0]['specs']
            else:
                equipment_display = ", ".join(all_items_df['equipment'].tolist())
                specs_display = "Multiple Items"
            qty_display = ", ".join(f"{v:g}" for v in all_items_df['quantity'].tolist())
            actual_price_display = ", ".join(f"{v:.0f}" for v in all_items_df['unit_actual_cost'].tolist())
            other_expenses_display = ", ".join(f"{v:.0f}" for v in all_items_df['other_expenses'].tolist())

        old_row = st.session_state.business_df[st.session_state.business_df['id'] == deal_id].iloc[0]
        paid = old_row['paid']
        remaining = close_deal - paid
        profit = close_deal - actual_cost
        status = "Paid" if remaining <= 0.01 else "Pending"

        cur.execute("""UPDATE business_deals SET equipment=?, specs=?, qty_per_item=?, close_deal=?,
                       actual_cost=?, actual_price_per_item=?, other_expenses_per_item=?,
                       remaining=?, profit=?, status=? WHERE id=?""",
                    (equipment_display, specs_display, qty_display, close_deal, actual_cost,
                     actual_price_display, other_expenses_display, remaining, profit, status, deal_id))
        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()
        st.session_state.business_df = read_sql_df("SELECT * FROM business_deals", conn)
        conn.close()

    with st.container(border=True):
        c1, c2 = st.columns(2)
        client = c1.text_input("Client Name/Hospital", key="deal_client")
        team_member = c2.text_input("Team Member (Optional)", key="deal_team_member")

        c3, c4 = st.columns(2)
        item_name = c3.text_input("Equipment Name", key="item_name_input")
        item_specs = c4.text_input("Specs", key="item_specs_input")

        c5, c6, c7, c8 = st.columns(4)
        item_qty = c5.number_input("Qty", min_value=1, format="%g", key="item_qty_input")
        item_price = c6.number_input("Unit Price", min_value=0.0, format="%g", key="item_price_input")
        item_cost = c7.number_input("Actual Cost per Unit", min_value=0.0, format="%g", key="item_cost_input")
        item_other = c8.number_input("Other Expenses", min_value=0.0, format="%g", key="item_other_input")

        paid = st.number_input("Payment sent by Client", min_value=0.0, format="%g", key="deal_paid")

        st.button("➕ Add to List", use_container_width=True, on_click=_add_item_cb)
        if st.session_state.get("add_item_warning"):
            st.warning("Equipment Name is required.")

        if st.session_state.temp_items:
            st.write("**Added Items:**")
            for idx, item in enumerate(st.session_state.temp_items):
                ic1, ic2, ic3, ic4, ic5, ic6, ic7 = st.columns([2, 2, 1, 1, 1, 1, 0.6])
                ic1.write(item['equipment']); ic2.write(item['specs'])
                ic3.write(f"{item['quantity']:g}"); ic4.write(f"{item['unit_price']:.0f}")
                ic5.write(f"{item['other_expenses']:.0f}")
                ic6.write(f"{item['line_total']:.0f}")
                ic7.button("🗑️", key=f"del_item_{idx}", on_click=_remove_item_cb, args=(idx,))
            st.caption(f"Running Total: Rs {sum(i['line_total'] for i in st.session_state.temp_items):,.0f}")

    with st.form("deal_form", clear_on_submit=True):
        st.form_submit_button("✅ Log Deal", use_container_width=True, on_click=_log_deal_cb)

    if st.session_state.get("deal_message"):
        level, text = st.session_state.deal_message
        getattr(st, level)(text)
        st.session_state.deal_message = None

    st.divider()
    st.subheader("📋 Records")

    _all_dates = pd.to_datetime(st.session_state.business_df['date'], errors='coerce') \
        if not st.session_state.business_df.empty else pd.Series([], dtype='datetime64[ns]')
    _default_from = _all_dates.min().date() if not _all_dates.empty and _all_dates.notna().any() else date.today()
    _default_to = _all_dates.max().date() if not _all_dates.empty and _all_dates.notna().any() else date.today()
    rf1, rf2 = st.columns(2)
    records_from = rf1.date_input("From", value=_default_from, key="records_filter_from")
    records_to = rf2.date_input("To", value=_default_to, key="records_filter_to")

    display_df = st.session_state.business_df.copy()
    for col in RECORD_DISPLAY_COLUMNS:
        if col not in display_df.columns:
            display_df[col] = None
    display_df = display_df[RECORD_DISPLAY_COLUMNS]
    display_df = _apply_date_filter(display_df, 'date', records_from, records_to)

    # Status is a manual dropdown (Decline / Approved / Paid / Pending) edited
    # directly in the table below. It is no longer auto-overwritten from
    # Remaining -- whatever was last saved in the DB is shown as-is; any
    # legacy/blank value falls back to "Pending" so the dropdown always has
    # a valid selection.
    if not display_df.empty:
        display_df['status'] = display_df['status'].apply(
            lambda s: s if s in STATUS_OPTIONS else 'Pending')

    edited_records = st.data_editor(
        display_df, use_container_width=True, hide_index=True, num_rows="fixed",
        disabled=["id"], key="records_editor_data",
        column_config={
            "status": st.column_config.SelectboxColumn("status", options=STATUS_OPTIONS, required=True)
        }
    )
    if st.button("💾 Save Records Changes", key="save_records_btn"):
        _save_records_edits(edited_records)
        st.success("Records updated successfully!")

    if not edited_records.empty:
        # Live preview of what's currently in the editor. Only the Status
        # cell itself is highlighted (never the whole row) when it reads
        # "Paid" -- a subtle, high-contrast green so the text stays legible.
        preview_df = edited_records.copy()
        preview_view = preview_df[RECORD_DISPLAY_COLUMNS].copy()
        preview_view = _prepare_export_df(
            preview_view, date_cols=['date'],
            money_cols=['close_deal', 'actual_cost', 'paid', 'remaining', 'profit'])
        preview_view = preview_view.rename(columns=dict(zip(RECORD_DISPLAY_COLUMNS, RECORD_DISPLAY_HEADERS)))

        st.caption("🟢 The Status cell only (not the whole row) is highlighted when it reads \"Paid\".")
        st.dataframe(preview_view.style.map(_highlight_paid_status_cell, subset=['Status']),
             width='stretch', hide_index=True)

    if not display_df.empty:
        export_df = display_df.sort_values('id', ascending=False)
        export_df = _prepare_export_df(
            export_df, date_cols=['date'],
            money_cols=['close_deal', 'actual_cost', 'paid', 'remaining', 'profit'])
        export_df_named = export_df.rename(columns=dict(zip(RECORD_DISPLAY_COLUMNS, RECORD_DISPLAY_HEADERS)))
        rc1, rc2 = st.columns(2)
        rc1.download_button("⬇️ Records CSV", data=_df_to_csv_bytes(export_df_named),
                             file_name="records.csv", mime="text/csv", key="records_csv_btn")
        records_pdf_path = generate_sheet_pdf(export_df, RECORD_DISPLAY_HEADERS, RECORD_COL_WIDTHS,
                                               "Records", "records_sheet", orientation="L")
        with open(records_pdf_path, "rb") as f:
            rc2.download_button("⬇️ Records PDF", data=f, file_name="records.pdf",
                                 mime="application/pdf", key="records_pdf_btn")

    # =====================================================================
    # Manage Deals -- Edit / Correct / Add Items / Delete Whole Deal
    # =====================================================================
    st.divider()
    st.subheader("🔧 Manage Deals (Add/Edit/Delete Items)")
    st.caption("➕ Manage = add new items. ✏️ Correct = fix wrong equipment/specs/qty/price "
               "or delete an item. 🗑️ Delete = permanently delete the entire deal.")

    if st.session_state.get("deal_delete_message"):
        level, text = st.session_state.deal_delete_message
        getattr(st, level)(text)
        st.session_state.deal_delete_message = None

    if not st.session_state.business_df.empty:

        def _select_deal_to_edit_cb(deal_id):
            st.session_state.editing_deal_id = deal_id
            st.session_state.update_temp_items = []

        header_c1, header_c2, header_c3, header_c4, header_c5 = st.columns([1, 3, 2, 1, 1])
        header_c1.markdown("**No.**"); header_c2.markdown("**Client**")
        header_c3.markdown("**Close Deal**"); header_c4.markdown("**Manage**"); header_c5.markdown("**Delete**")

        for _, drow in st.session_state.business_df.sort_values('id', ascending=False).iterrows():
            d_id = int(drow['id'])
            rcol1, rcol2, rcol3, rcol4, rcol5 = st.columns([1, 3, 2, 1, 1])
            rcol1.write(f"#{d_id}")
            rcol2.write(drow['client'])
            rcol3.write(f"Rs {drow['close_deal']:,.0f}")
            rcol4.button("✏️ Manage", key=f"edit_deal_btn_{d_id}",
                        on_click=_select_deal_to_edit_cb, args=(d_id,))

            if st.session_state.get("confirm_delete_deal_id") == d_id:
                dcol_a, dcol_b = rcol5.columns(2)
                dcol_a.button("✅", key=f"confirm_del_{d_id}", on_click=_delete_deal_cb, args=(d_id,),
                              type="primary", help="Confirm delete")
                dcol_b.button("✖️", key=f"cancel_del_{d_id}", on_click=_cancel_delete_deal_cb,
                              help="Cancel")
            else:
                rcol5.button("🗑️", key=f"del_deal_btn_{d_id}", on_click=_request_delete_deal_cb, args=(d_id,))

        if st.session_state.get("editing_deal_id") is not None:
            edit_deal_id = st.session_state.editing_deal_id

            with st.container(border=True):
                st.markdown(f"**Managing Deal #{edit_deal_id}:**")

                conn = get_connection()
                existing_items_df = read_sql_df(
                    "SELECT id, equipment, specs, quantity, unit_price, unit_actual_cost, other_expenses, line_total "
                    "FROM deal_items WHERE deal_id = ?", conn, params=(int(edit_deal_id),))
                conn.close()

                st.write("**✏️ Existing Items (correct mistakes here, or delete a row):**")
                edited_existing_items = st.data_editor(
                    existing_items_df, use_container_width=True, hide_index=True,
                    num_rows="dynamic", disabled=["id", "line_total"],
                    key=f"existing_items_editor_{edit_deal_id}"
                )
                if st.button("💾 Save Correction", key=f"save_existing_items_btn_{edit_deal_id}"):
                    _save_existing_items_cb(edit_deal_id, edited_existing_items)
                    st.success(f"Deal #{edit_deal_id} items updated!")
                    st.rerun()

                st.divider()
                st.write("**➕ Add a brand new item to this deal:**")

                uc3, uc4 = st.columns(2)
                uc3.text_input("Equipment Name", key="update_item_name_input")
                uc4.text_input("Specs", key="update_item_specs_input")
                uc5, uc6, uc7, uc8 = st.columns(4)
                uc5.number_input("Qty", min_value=1, format="%g", key="update_item_qty_input")
                uc6.number_input("Unit Price", min_value=0.0, format="%g", key="update_item_price_input")
                uc7.number_input("Actual Cost per Unit", min_value=0.0, format="%g", key="update_item_cost_input")
                uc8.number_input("Other Expenses", min_value=0.0, format="%g", key="update_item_other_input")

                def _add_update_item_cb():
                    name = st.session_state.update_item_name_input
                    if name and name.strip():
                        qty = st.session_state.update_item_qty_input
                        price = st.session_state.update_item_price_input
                        cost = st.session_state.update_item_cost_input
                        other = st.session_state.update_item_other_input
                        st.session_state.update_temp_items.append({
                            'equipment': name,
                            'specs': st.session_state.update_item_specs_input,
                            'quantity': qty,
                            'unit_price': price,
                            'unit_actual_cost': cost,
                            'other_expenses': other,
                            'line_total': qty * price,
                            'line_actual_cost': qty * cost + other,
                        })
                        st.session_state.update_item_name_input = ""
                        st.session_state.update_item_specs_input = ""
                        st.session_state.update_item_qty_input = 1
                        st.session_state.update_item_price_input = 0.0
                        st.session_state.update_item_cost_input = 0.0
                        st.session_state.update_item_other_input = 0.0
                        st.session_state.update_add_item_warning = False
                    else:
                        st.session_state.update_add_item_warning = True

                st.button("➕ Add to List", key="update_add_item_btn", on_click=_add_update_item_cb)
                if st.session_state.get("update_add_item_warning"):
                    st.warning("Equipment Name is required.")

                def _remove_update_item_cb(idx):
                    if 0 <= idx < len(st.session_state.update_temp_items):
                        st.session_state.update_temp_items.pop(idx)

                if st.session_state.update_temp_items:
                    st.write("**New Items (not yet saved):**")
                    for idx, item in enumerate(st.session_state.update_temp_items):
                        nic1, nic2, nic3, nic4, nic5, nic6, nic7 = st.columns([2, 2, 1, 1, 1, 1, 0.6])
                        nic1.write(item['equipment']); nic2.write(item['specs'])
                        nic3.write(f"{item['quantity']:g}"); nic4.write(f"{item['unit_price']:.0f}")
                        nic5.write(f"{item['other_expenses']:.0f}")
                        nic6.write(f"{item['line_total']:.0f}")
                        nic7.button("🗑️", key=f"del_update_item_{idx}",
                                    on_click=_remove_update_item_cb, args=(idx,))

                    def _update_deal_cb():
                        deal_id = int(edit_deal_id)
                        new_items = st.session_state.update_temp_items
                        if not new_items:
                            return
                        conn = get_connection()
                        cur = conn.cursor()
                        for item in new_items:
                            cur.execute("""INSERT INTO deal_items
                                (deal_id, equipment, specs, quantity, unit_price, unit_actual_cost, other_expenses, line_total, line_actual_cost)
                                VALUES (?,?,?,?,?,?,?,?,?)""",
                                (deal_id, item['equipment'], item['specs'], item['quantity'], item['unit_price'],
                                 item['unit_actual_cost'], item['other_expenses'], item['line_total'], item['line_actual_cost']))
                        conn.commit()
                        if hasattr(conn, "sync"):
                            conn.sync()

                        all_items_df = read_sql_df("SELECT * FROM deal_items WHERE deal_id=?", conn, params=(deal_id,))
                        close_deal = all_items_df['line_total'].sum()
                        actual_cost = all_items_df['line_actual_cost'].sum()
                        if len(all_items_df) == 1:
                            equipment_display = all_items_df.iloc[0]['equipment']
                            specs_display = all_items_df.iloc[0]['specs']
                        else:
                            equipment_display = ", ".join(all_items_df['equipment'].tolist())
                            specs_display = "Multiple Items"
                        qty_display = ", ".join(f"{v:g}" for v in all_items_df['quantity'].tolist())
                        actual_price_display = ", ".join(f"{v:.0f}" for v in all_items_df['unit_actual_cost'].tolist())
                        other_expenses_display = ", ".join(f"{v:.0f}" for v in all_items_df['other_expenses'].tolist())

                        old_row = st.session_state.business_df[st.session_state.business_df['id'] == deal_id].iloc[0]
                        paid = old_row['paid']
                        remaining = close_deal - paid
                        profit = close_deal - actual_cost
                        status = "Paid" if remaining <= 0.01 else "Pending"

                        cur.execute("""UPDATE business_deals SET equipment=?, specs=?, qty_per_item=?, close_deal=?,
                                       actual_cost=?, actual_price_per_item=?, other_expenses_per_item=?,
                                       remaining=?, profit=?, status=? WHERE id=?""",
                                    (equipment_display, specs_display, qty_display, close_deal, actual_cost,
                                     actual_price_display, other_expenses_display, remaining, profit, status, deal_id))
                        conn.commit()
                        if hasattr(conn, "sync"):
                            conn.sync()
                        st.session_state.business_df = read_sql_df("SELECT * FROM business_deals", conn)
                        conn.close()

                        st.session_state.update_temp_items = []
                        st.session_state.editing_deal_id = None
                        st.session_state.update_deal_message = (
                            "success", f"Deal #{deal_id} updated -- new items added! "
                                       f"The Invoice/Quotation/Challan for this deal will now print with the updated items.")

                    st.button("💾 Update Deal (Save New Items)", key="update_deal_btn",
                             on_click=_update_deal_cb, type="primary")

                def _cancel_edit_deal_cb():
                    st.session_state.editing_deal_id = None
                    st.session_state.update_temp_items = []

                st.button("✖️ Close", key="cancel_edit_deal_btn", on_click=_cancel_edit_deal_cb)

        if st.session_state.get("update_deal_message"):
            level, text = st.session_state.update_deal_message
            getattr(st, level)(text)
            st.session_state.update_deal_message = None

    st.divider()
    st.subheader("🖨️ Generate Invoice / Quotation / Delivery Challan")
    if not st.session_state.business_df.empty:
        col_a, col_b = st.columns([0.6, 0.4])
        selected_id = col_a.selectbox("Select Deal ID:", st.session_state.business_df['id'].tolist())
        doc_choice = col_b.selectbox("Print as", ["Invoice", "Quotation", "Delivery Challan"])

        terms_text = None
        if doc_choice == "Quotation":
            terms_text = st.text_area(
                "Terms & Conditions (printed on the Quotation -- editable)",
                "1. 50% advance required, remaining on delivery.\n"
                "2. Prices are valid for 15 days from the quotation date.\n"
                "3. Delivery within 7-10 working days after confirmation."
            )

        conn = get_connection()
        deal_row = st.session_state.business_df[st.session_state.business_df['id'] == selected_id].iloc[0]
        items_df = read_sql_df("SELECT * FROM deal_items WHERE deal_id = ?", conn, params=(int(selected_id),))
        conn.close()

        dl1, dl2 = st.columns(2)

        pdf_path = generate_pdf(deal_row, items_df, doc_type=doc_choice, terms_text=terms_text)
        with open(pdf_path, "rb") as f:
            dl1.download_button(f"📥 Download {doc_choice} (PDF)", f, file_name=pdf_path,
                                 mime="application/pdf", use_container_width=True)

        docx_path = generate_docx(deal_row, items_df, doc_type=doc_choice, terms_text=terms_text)
        with open(docx_path, "rb") as f:
            dl2.download_button(f"📄 Download {doc_choice} (Word)", f, file_name=docx_path,
                                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                 use_container_width=True)
    else:
        st.info("No records yet.")

# ---------------- TAB 3: CREDIT / DEBIT / EXPENSE SHEETS ----------------
with tab3:
    st.title("💳 Credit / Debit / Expense Sheets")
    st.caption(
        "These sheets update automatically from deal records. If a client has paid in full or "
        "partially, that deal appears in the Credit Sheet (the client still owes us). "
        "If a client has overpaid, that deal appears in the Debit Sheet "
        "(we owe the client a refund). You can also add manual entries below "
        "(amounts owed outside of deals)."
    )

    def _add_credit_cb():
        name = st.session_state.credit_new_client
        if name and name.strip():
            total_payment = st.session_state.credit_new_total
            paid_by = st.session_state.credit_new_paid
            remaining = total_payment - paid_by
            conn = get_connection()
            conn.execute(
                "INSERT INTO credit_manual (client, total_payment, paid_by_client, remaining_from_client) VALUES (?,?,?,?)",
                (name, total_payment, paid_by, remaining))
            conn.commit()
            if hasattr(conn, "sync"):
                conn.sync()
            st.session_state.credit_manual_df = read_sql_df("SELECT * FROM credit_manual", conn)
            conn.close()
            st.session_state.credit_new_client = ""
            st.session_state.credit_new_total = 0.0
            st.session_state.credit_new_paid = 0.0
            st.session_state.credit_add_warning = False
        else:
            st.session_state.credit_add_warning = True

    def _add_debit_cb():
        name = st.session_state.debit_new_client
        if name and name.strip():
            total_payment = st.session_state.debit_new_total
            paid_to = st.session_state.debit_new_paid
            remaining = total_payment - paid_to
            conn = get_connection()
            conn.execute(
                "INSERT INTO debit_manual (client, total_payment, paid_to_client, remaining_to_be_paid) VALUES (?,?,?,?)",
                (name, total_payment, paid_to, remaining))
            conn.commit()
            if hasattr(conn, "sync"):
                conn.sync()
            st.session_state.debit_manual_df = read_sql_df("SELECT * FROM debit_manual", conn)
            conn.close()
            st.session_state.debit_new_client = ""
            st.session_state.debit_new_total = 0.0
            st.session_state.debit_new_paid = 0.0
            st.session_state.debit_add_warning = False
        else:
            st.session_state.debit_add_warning = True

    def _save_credit_edits(edited):
        conn = get_connection()
        conn.execute("DELETE FROM credit_manual")
        for row in edited.to_dict("records"):
            if str(row.get('client', '')).strip():
                total_payment = row.get('total_payment', 0) or 0
                paid_by = row.get('paid_by_client', 0) or 0
                conn.execute(
                    "INSERT INTO credit_manual (client, total_payment, paid_by_client, remaining_from_client) VALUES (?,?,?,?)",
                    (row['client'], total_payment, paid_by, total_payment - paid_by))
        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()
        st.session_state.credit_manual_df = read_sql_df("SELECT * FROM credit_manual", conn)
        conn.close()

    def _save_debit_edits(edited):
        conn = get_connection()
        conn.execute("DELETE FROM debit_manual")
        for row in edited.to_dict("records"):
            if str(row.get('client', '')).strip():
                total_payment = row.get('total_payment', 0) or 0
                paid_to = row.get('paid_to_client', 0) or 0
                conn.execute(
                    "INSERT INTO debit_manual (client, total_payment, paid_to_client, remaining_to_be_paid) VALUES (?,?,?,?)",
                    (row['client'], total_payment, paid_to, total_payment - paid_to))
        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()
        st.session_state.debit_manual_df = read_sql_df("SELECT * FROM debit_manual", conn)
        conn.close()

    def _add_expense_cb():
        category = st.session_state.expense_category_input
        description = (st.session_state.expense_desc_manual_input
                        if category == "Others" else category)
        if not description or not str(description).strip():
            st.session_state.expense_add_warning = True
            return
        amount = st.session_state.expense_amount_input
        exp_date = st.session_state.expense_date_input
        conn = get_connection()
        conn.execute(
            "INSERT INTO daily_expenses (date, category, description, amount) VALUES (?,?,?,?)",
            (str(exp_date), category, description, amount))
        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()
        st.session_state.expense_df = read_sql_df("SELECT * FROM daily_expenses", conn)
        conn.close()
        st.session_state.expense_desc_manual_input = ""
        st.session_state.expense_amount_input = 0.0
        st.session_state.expense_add_warning = False

    def _save_expense_edits(edited):
        conn = get_connection()
        conn.execute("DELETE FROM daily_expenses")
        for row in edited.to_dict("records"):
            desc = row.get('description', '')
            if str(desc).strip():
                conn.execute(
                    "INSERT INTO daily_expenses (date, category, description, amount) VALUES (?,?,?,?)",
                    (row.get('date'), row.get('category'), desc, row.get('amount', 0) or 0))
        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()
        st.session_state.expense_df = read_sql_df("SELECT * FROM daily_expenses", conn)
        conn.close()

    deals = st.session_state.business_df
    credit_tab, debit_tab, expense_tab = st.tabs([
        "💰 Credit Sheet (Client owes us)",
        "💸 Debit Sheet (We owe client)",
        "🧾 Daily Expense Sheet"
    ])

    # ---------------- CREDIT SHEET ----------------
    with credit_tab:
        _credit_dates = pd.to_datetime(deals['date'], errors='coerce') if not deals.empty else pd.Series([], dtype='datetime64[ns]')
        _credit_default_from = _credit_dates.min().date() if not _credit_dates.empty and _credit_dates.notna().any() else date.today()
        _credit_default_to = _credit_dates.max().date() if not _credit_dates.empty and _credit_dates.notna().any() else date.today()
        cf1, cf2 = st.columns(2)
        credit_from = cf1.date_input("From", value=_credit_default_from, key="credit_filter_from")
        credit_to = cf2.date_input("To", value=_credit_default_to, key="credit_filter_to")
        st.caption("The filter only applies to deal entries -- manual entries have no date and will always be shown.")

        auto_credit = deals[deals['remaining'] >= 0].copy() if not deals.empty else deals
        auto_credit = _apply_date_filter(auto_credit, 'date', credit_from, credit_to)
        auto_credit_view = pd.DataFrame({
            "Client Name": auto_credit['client'],
            "Id No": "D-" + auto_credit['id'].astype(str),
            "Total Payment": auto_credit['close_deal'],
            "Paid by Client": auto_credit['paid'],
            "Remaining from Client": auto_credit['remaining'],
        }) if not auto_credit.empty else pd.DataFrame(
            columns=["Client Name", "Id No", "Total Payment", "Paid by Client", "Remaining from Client"])

        manual_c = st.session_state.credit_manual_df
        manual_c_view = pd.DataFrame({
            "Client Name": manual_c['client'],
            "Id No": "C-" + manual_c['id'].astype(str),
            "Total Payment": manual_c['total_payment'],
            "Paid by Client": manual_c['paid_by_client'],
            "Remaining from Client": manual_c['remaining_from_client'],
        }) if not manual_c.empty else pd.DataFrame(
            columns=["Client Name", "Id No", "Total Payment", "Paid by Client", "Remaining from Client"])

        full_credit_view = pd.concat([auto_credit_view, manual_c_view], ignore_index=True)

        st.subheader("📋 Full Credit Sheet")
        st.dataframe(full_credit_view, use_container_width=True, hide_index=True)

        if not full_credit_view.empty:
            full_credit_export = _prepare_export_df(
                full_credit_view,
                money_cols=["Total Payment", "Paid by Client", "Remaining from Client"])
            cec1, cec2 = st.columns(2)
            cec1.download_button("⬇️ Credit Sheet CSV", data=_df_to_csv_bytes(full_credit_export),
                                  file_name="credit_sheet.csv", mime="text/csv", key="credit_csv_btn")
            credit_pdf_path = generate_sheet_pdf(full_credit_export, CREDIT_HEADERS, CREDIT_COL_WIDTHS,
                                                  "Credit Sheet", "credit_sheet", orientation="L")
            with open(credit_pdf_path, "rb") as f:
                cec2.download_button("⬇️ Credit Sheet PDF", data=f, file_name="credit_sheet.pdf",
                                      mime="application/pdf", key="credit_pdf_btn")

        st.divider()
        st.subheader("➕ Add Manual Entry (outside deals, amounts receivable)")
        cc1, cc2, cc3 = st.columns(3)
        cc1.text_input("Client Name", key="credit_new_client")
        cc2.number_input("Total Payment", min_value=0.0, format="%g", key="credit_new_total")
        cc3.number_input("Paid by Client", min_value=0.0, format="%g", key="credit_new_paid")
        st.button("➕ Add to Credit Sheet", on_click=_add_credit_cb, key="add_credit_btn")
        if st.session_state.get("credit_add_warning"):
            st.warning("Client Name is required.")

        if not st.session_state.credit_manual_df.empty:
            st.write("**Manual Entries (editable)**")
            edited_credit = st.data_editor(
                st.session_state.credit_manual_df.drop(columns=['id']),
                use_container_width=True, hide_index=True, num_rows="dynamic", key="credit_editor_data")
            if st.button("💾 Save Credit Changes", key="save_credit_btn"):
                _save_credit_edits(edited_credit)
                st.success("Credit Sheet updated!")

    # ---------------- DEBIT SHEET ----------------
    with debit_tab:
        _debit_dates = pd.to_datetime(deals['date'], errors='coerce') if not deals.empty else pd.Series([], dtype='datetime64[ns]')
        _debit_default_from = _debit_dates.min().date() if not _debit_dates.empty and _debit_dates.notna().any() else date.today()
        _debit_default_to = _debit_dates.max().date() if not _debit_dates.empty and _debit_dates.notna().any() else date.today()
        df1, df2 = st.columns(2)
        debit_from = df1.date_input("From", value=_debit_default_from, key="debit_filter_from")
        debit_to = df2.date_input("To", value=_debit_default_to, key="debit_filter_to")
        st.caption("The filter only applies to deal entries -- manual entries have no date and will always be shown.")

        auto_debit = deals[deals['remaining'] < 0].copy() if not deals.empty else deals
        auto_debit = _apply_date_filter(auto_debit, 'date', debit_from, debit_to)
        auto_debit_view = pd.DataFrame({
            "Client Name": auto_debit['client'],
            "Id No": "D-" + auto_debit['id'].astype(str),
            "Total Payment": auto_debit['close_deal'],
            "Paid to Client": 0,
            "Remaining to be paid": auto_debit['remaining'].abs(),
        }) if not auto_debit.empty else pd.DataFrame(
            columns=["Client Name", "Id No", "Total Payment", "Paid to Client", "Remaining to be paid"])

        manual_d = st.session_state.debit_manual_df
        manual_d_view = pd.DataFrame({
            "Client Name": manual_d['client'],
            "Id No": "C-" + manual_d['id'].astype(str),
            "Total Payment": manual_d['total_payment'],
            "Paid to Client": manual_d['paid_to_client'],
            "Remaining to be paid": manual_d['remaining_to_be_paid'],
        }) if not manual_d.empty else pd.DataFrame(
            columns=["Client Name", "Id No", "Total Payment", "Paid to Client", "Remaining to be paid"])

        full_debit_view = pd.concat([auto_debit_view, manual_d_view], ignore_index=True)

        st.subheader("📋 Full Debit Sheet")
        st.dataframe(full_debit_view, use_container_width=True, hide_index=True)

        if not full_debit_view.empty:
            full_debit_export = _prepare_export_df(
                full_debit_view,
                money_cols=["Total Payment", "Paid to Client", "Remaining to be paid"])
            dec1, dec2 = st.columns(2)
            dec1.download_button("⬇️ Debit Sheet CSV", data=_df_to_csv_bytes(full_debit_export),
                                  file_name="debit_sheet.csv", mime="text/csv", key="debit_csv_btn")
            debit_pdf_path = generate_sheet_pdf(full_debit_export, DEBIT_HEADERS, DEBIT_COL_WIDTHS,
                                                 "Debit Sheet", "debit_sheet", orientation="L")
            with open(debit_pdf_path, "rb") as f:
                dec2.download_button("⬇️ Debit Sheet PDF", data=f, file_name="debit_sheet.pdf",
                                      mime="application/pdf", key="debit_pdf_btn")

        st.divider()
        st.subheader("➕ Add Manual Entry (outside deals, amounts payable)")
        dc1, dc2, dc3 = st.columns(3)
        dc1.text_input("Client Name", key="debit_new_client")
        dc2.number_input("Total Payment", min_value=0.0, format="%g", key="debit_new_total")
        dc3.number_input("Paid to Client", min_value=0.0, format="%g", key="debit_new_paid")
        st.button("➕ Add to Debit Sheet", on_click=_add_debit_cb, key="add_debit_btn")
        if st.session_state.get("debit_add_warning"):
            st.warning("Client Name is required.")

        if not st.session_state.debit_manual_df.empty:
            st.write("**Manual Entries (editable)**")
            edited_debit = st.data_editor(
                st.session_state.debit_manual_df.drop(columns=['id']),
                use_container_width=True, hide_index=True, num_rows="dynamic", key="debit_editor_data")
            if st.button("💾 Save Debit Changes", key="save_debit_btn"):
                _save_debit_edits(edited_debit)
                st.success("Debit Sheet updated!")

    # ---------------- DAILY EXPENSE SHEET ----------------
    with expense_tab:
        st.subheader("🧾 Daily Expense Sheet")
        st.caption("Selecting Eating/Fuel auto-fills the description. "
                    "Selecting Others requires you to type your own description.")

        ec1, ec2, ec3, ec4 = st.columns([1.2, 1.6, 1, 1])
        category = ec1.selectbox("Category", ["Eating", "Fuel", "Others"], key="expense_category_input")
        if category == "Others":
            ec2.text_input("Description", key="expense_desc_manual_input")
        else:
            ec2.text_input("Description", value=category, disabled=True)
        ec3.number_input("Amount", min_value=0.0, format="%g", key="expense_amount_input")
        ec4.date_input("Date", value=date.today(), key="expense_date_input")

        st.button("➕ Add Expense", on_click=_add_expense_cb, key="add_expense_btn")
        if st.session_state.get("expense_add_warning"):
            st.warning("Description is required.")

        st.divider()
        st.subheader("📋 Full Expense Sheet")

        _exp_dates = pd.to_datetime(st.session_state.expense_df['date'], errors='coerce') \
            if not st.session_state.expense_df.empty else pd.Series([], dtype='datetime64[ns]')
        _exp_default_from = _exp_dates.min().date() if not _exp_dates.empty and _exp_dates.notna().any() else date.today()
        _exp_default_to = _exp_dates.max().date() if not _exp_dates.empty and _exp_dates.notna().any() else date.today()
        ef1, ef2 = st.columns(2)
        expense_from = ef1.date_input("From", value=_exp_default_from, key="expense_filter_from")
        expense_to = ef2.date_input("To", value=_exp_default_to, key="expense_filter_to")

        filtered_expense_df = _apply_date_filter(
            st.session_state.expense_df.drop(columns=['id'], errors='ignore'), 'date', expense_from, expense_to)

        st.dataframe(filtered_expense_df, use_container_width=True, hide_index=True)

        _expense_total = filtered_expense_df['amount'].sum() if not filtered_expense_df.empty else 0
        st.metric("💵 Total Expense (selected range)", f"Rs {_expense_total:,.0f}")

        if not filtered_expense_df.empty:
            expense_export = _prepare_export_df(filtered_expense_df, date_cols=['date'], money_cols=['amount'])
            eec1, eec2 = st.columns(2)
            eec1.download_button("⬇️ Expense Sheet CSV",
                                  data=_df_to_csv_bytes(expense_export),
                                  file_name="expense_sheet.csv", mime="text/csv", key="expense_csv_btn")
            expense_pdf_df = expense_export[['description', 'amount']].rename(
                columns={'description': 'Description', 'amount': 'Amount'})
            expense_pdf_path = generate_sheet_pdf(expense_pdf_df, EXPENSE_HEADERS, EXPENSE_COL_WIDTHS,
                                                   "Expense Sheet", "expense_sheet", orientation="P")
            with open(expense_pdf_path, "rb") as f:
                eec2.download_button("⬇️ Expense Sheet PDF", data=f, file_name="expense_sheet.pdf",
                                      mime="application/pdf", key="expense_pdf_btn")

        if not st.session_state.expense_df.empty:
            st.write("**Manual Entries (editable)**")
            edited_expense = st.data_editor(
                st.session_state.expense_df.drop(columns=['id']),
                use_container_width=True, hide_index=True, num_rows="dynamic", key="expense_editor_data",
                column_config={
                    "category": st.column_config.SelectboxColumn("category", options=["Eating", "Fuel", "Others"])
                })
            if st.button("💾 Save Expense Changes", key="save_expense_btn"):
                _save_expense_edits(edited_expense)
                st.success("Expense Sheet updated!")

# ---------------- TAB 4: PERFORMANCE INSIGHTS ----------------
with tab4:
    st.title("📊 Performance Insights")

    this_month = datetime.now().strftime("%Y-%m")
    biz = st.session_state.business_df.copy()
    if not biz.empty:
        biz_dates = pd.to_datetime(biz['date'], errors='coerce')
        monthly_paid = biz.loc[biz_dates.dt.strftime("%Y-%m") == this_month, 'paid'].sum()
        pending_total = biz.loc[biz['remaining'] > 0, 'remaining'].sum()
        total_profit = biz['profit'].sum()
    else:
        monthly_paid = pending_total = total_profit = 0

    exp = st.session_state.expense_df.copy()
    if not exp.empty:
        exp_dates = pd.to_datetime(exp['date'], errors='coerce')
        monthly_expense = exp.loc[exp_dates.dt.strftime("%Y-%m") == this_month, 'amount'].sum()
    else:
        monthly_expense = 0

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("💰 Payments Received (This Month)", f"Rs {int(monthly_paid):,}")
    m2.metric("⏳ Pending Payments", f"Rs {int(pending_total):,}")
    m3.metric("📈 Total Profit", f"Rs {int(total_profit):,}")
    m4.metric("💸 Expenses (This Month)", f"Rs {int(monthly_expense):,}")

    st.divider()

    if not st.session_state.business_df.empty:
        st.metric("Total Revenue", f"Rs {int(st.session_state.business_df['close_deal'].sum()):,}")
        fig = px.bar(st.session_state.business_df, x='id', y='close_deal', template="plotly_dark")
        st.plotly_chart(fig, use_container_width=True)

# ---------------- TAB 5: LIABILITIES ----------------
with tab5:
    st.title("📉 Liabilities")
    st.caption("Track amounts your business owes -- loans, supplier dues, pending payables, etc.")

    def _add_liability_cb():
        desc = st.session_state.liability_desc_input
        if desc and desc.strip():
            total = st.session_state.liability_total_input
            paid = st.session_state.liability_paid_input
            remaining = total - paid
            liab_date = st.session_state.liability_date_input
            conn = get_connection()
            conn.execute(
                "INSERT INTO liabilities (date, description, total_amount, paid_amount, remaining) VALUES (?,?,?,?,?)",
                (str(liab_date), desc, total, paid, remaining))
            conn.commit()
            if hasattr(conn, "sync"):
                conn.sync()
            st.session_state.liability_df = read_sql_df("SELECT * FROM liabilities", conn)
            conn.close()
            st.session_state.liability_desc_input = ""
            st.session_state.liability_total_input = 0.0
            st.session_state.liability_paid_input = 0.0
            st.session_state.liability_add_warning = False
        else:
            st.session_state.liability_add_warning = True

    def _save_liability_edits(edited):
        conn = get_connection()
        conn.execute("DELETE FROM liabilities")
        for row in edited.to_dict("records"):
            if str(row.get('description', '')).strip():
                total = row.get('total_amount', 0) or 0
                paid = row.get('paid_amount', 0) or 0
                conn.execute(
                    "INSERT INTO liabilities (date, description, total_amount, paid_amount, remaining) VALUES (?,?,?,?,?)",
                    (row.get('date'), row.get('description'), total, paid, total - paid))
        conn.commit()
        if hasattr(conn, "sync"):
            conn.sync()
        st.session_state.liability_df = read_sql_df("SELECT * FROM liabilities", conn)
        conn.close()

    lc1, lc2, lc3, lc4 = st.columns([1, 2, 1, 1])
    lc1.date_input("Date", value=date.today(), key="liability_date_input")
    lc2.text_input("Description (e.g. Bank Loan, Supplier Due)", key="liability_desc_input")
    lc3.number_input("Total Amount", min_value=0.0, format="%g", key="liability_total_input")
    lc4.number_input("Paid Amount", min_value=0.0, format="%g", key="liability_paid_input")
    st.button("➕ Add Liability", on_click=_add_liability_cb, key="add_liability_btn")
    if st.session_state.get("liability_add_warning"):
        st.warning("Description is required.")

    st.divider()
    st.subheader("📋 Full Liabilities Sheet")

    liab_df = st.session_state.liability_df
    _liab_dates = pd.to_datetime(liab_df['date'], errors='coerce') \
        if not liab_df.empty else pd.Series([], dtype='datetime64[ns]')
    _liab_default_from = _liab_dates.min().date() if not _liab_dates.empty and _liab_dates.notna().any() else date.today()
    _liab_default_to = _liab_dates.max().date() if not _liab_dates.empty and _liab_dates.notna().any() else date.today()
    lf1, lf2 = st.columns(2)
    liability_from = lf1.date_input("From", value=_liab_default_from, key="liability_filter_from")
    liability_to = lf2.date_input("To", value=_liab_default_to, key="liability_filter_to")

    filtered_liab_df = _apply_date_filter(
        liab_df.drop(columns=['id'], errors='ignore'), 'date', liability_from, liability_to)
    st.dataframe(filtered_liab_df, use_container_width=True, hide_index=True)

    _liab_total_remaining = filtered_liab_df['remaining'].sum() if not filtered_liab_df.empty else 0
    st.metric("📉 Total Outstanding Liabilities (selected range)", f"Rs {_liab_total_remaining:,.0f}")

    if not filtered_liab_df.empty:
        liab_export = _prepare_export_df(filtered_liab_df, date_cols=['date'],
                                          money_cols=['total_amount', 'paid_amount', 'remaining'])
        liab_export_named = liab_export.rename(columns={
            'date': 'Date', 'description': 'Description', 'total_amount': 'Total Amount',
            'paid_amount': 'Paid Amount', 'remaining': 'Remaining'})
        lec1, lec2 = st.columns(2)
        lec1.download_button("⬇️ Liabilities CSV", data=_df_to_csv_bytes(liab_export_named),
                              file_name="liabilities.csv", mime="text/csv", key="liability_csv_btn")
        liab_pdf_df = liab_export_named[LIABILITY_HEADERS]
        liability_pdf_path = generate_sheet_pdf(liab_pdf_df, LIABILITY_HEADERS, LIABILITY_COL_WIDTHS,
                                                  "Liabilities Sheet", "liabilities_sheet", orientation="P")
        with open(liability_pdf_path, "rb") as f:
            lec2.download_button("⬇️ Liabilities PDF", data=f, file_name="liabilities.pdf",
                                  mime="application/pdf", key="liability_pdf_btn")

    if not st.session_state.liability_df.empty:
        st.write("**Manual Entries (editable)**")
        edited_liability = st.data_editor(
            st.session_state.liability_df.drop(columns=['id']),
            use_container_width=True, hide_index=True, num_rows="dynamic", key="liability_editor_data")
        if st.button("💾 Save Liabilities Changes", key="save_liability_btn"):
            _save_liability_edits(edited_liability)
            st.success("Liabilities updated!")

# ---------------- TAB 6: APPROVED (dedicated grouping for Approved status) ----------------
with tab6:
    st.title("✅ Approved")
    st.caption(
        "Entries whose Status is set to \"Approved\" (in the Business Deals → Records table) "
        "are grouped here automatically. Once a payment is processed for an approved entry, "
        "mark it as Paid right here -- its Status cell will then be highlighted."
    )

    approved_df = st.session_state.business_df[
        st.session_state.business_df['status'] == 'Approved'
    ].copy() if not st.session_state.business_df.empty else st.session_state.business_df

    if approved_df.empty:
        st.info("No entries are currently marked \"Approved\". Set a Record's Status dropdown "
                "to \"Approved\" in the Business Deals tab to see it grouped here.")
    else:
        st.subheader(f"📋 Approved Entries ({len(approved_df)})")

        for _, arow in approved_df.sort_values('id', ascending=False).iterrows():
            a_id = int(arow['id'])
            with st.container(border=True):
                ac1, ac2, ac3, ac4 = st.columns([2.5, 1.5, 1.5, 1.5])
                ac1.markdown(f"**#{a_id} — {arow['client']}**")
                ac2.write(f"Close Deal: Rs {arow['close_deal']:,.0f}")
                ac3.write(f"Paid: Rs {arow['paid']:,.0f}")
                ac4.write(f"Remaining: Rs {arow['remaining']:,.0f}")

                pay_col, decline_col = st.columns(2)
                pay_col.button(
                    "💰 Process Payment & Mark as Paid", key=f"approved_mark_paid_{a_id}",
                    on_click=_set_deal_status_cb, args=(a_id, "Paid"),
                    use_container_width=True, type="primary"
                )
                decline_col.button(
                    "✖️ Decline Instead", key=f"approved_decline_{a_id}",
                    on_click=_set_deal_status_cb, args=(a_id, "Decline"),
                    use_container_width=True
                )

        st.divider()
        st.subheader("📋 Approved Sheet (with Paid entries highlighted)")

        approved_view_cols = ['id', 'date', 'client', 'close_deal', 'paid', 'remaining', 'status']
        approved_view_headers = ['No.', 'Date', 'Client', 'Close Deal', 'Paid', 'Remaining', 'Status']

        # Show current-DB entries that are Approved OR were just moved to Paid
        # from this tab, so the "mark as paid" transition stays visible here
        # (payment handling happens right in this dedicated view).
        approved_sheet_df = st.session_state.business_df[
            st.session_state.business_df['status'].isin(['Approved', 'Paid'])
        ].copy()
        approved_sheet_view = approved_sheet_df[approved_view_cols].sort_values('id', ascending=False)
        approved_sheet_view = _prepare_export_df(
            approved_sheet_view, date_cols=['date'], money_cols=['close_deal', 'paid', 'remaining'])
        approved_sheet_view = approved_sheet_view.rename(
            columns=dict(zip(approved_view_cols, approved_view_headers)))

        st.caption("🟢 The Status cell only (not the whole row) is highlighted when it reads \"Paid\".")
        st.dataframe(
            approved_sheet_view.style.map(_highlight_paid_status_cell, subset=['Status']),
            width='stretch', hide_index=True
        )
            
        
